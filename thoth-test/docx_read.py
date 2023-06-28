#!/usr/bin/python
# -*- coding: utf-8 -*-
#
# *****************************************************
#
# file:     word_read.py
# author:   zlw2008ok@126.com
# date:     2023/1/16
# desc:     
#
# cmd>e.g.:
# *****************************************************
"""
python-docx模块虽然强大，但却不能处理后缀为".doc"的word文件。需要先通过批量处理将.doc转化为.docx类型（windows安装pypiwin32，mac安装LibreOffice）
"""
import os
import time

## mac doc to docx
import subprocess
import docx
from ctypes.wintypes import PUINT
source = '/Users/zlw/Documents/17-临床路径/国家发布的1010个临床路径汇总/临床路径汇总/doc/县医院版'
dest = '/Users/zlw/Documents/17-临床路径/国家发布的1010个临床路径汇总/临床路径汇总/docx'
g = os.listdir(source)
file_path = [f for f in g if f.endswith(('.doc'))]
print(file_path)
for i in file_path:
    print(i)
    file = (source + '/' + i )
    output = subprocess.check_output(["/Applications/LibreOffice.app/Contents/MacOS/soffice","--headless","--convert-to","docx",file,"--outdir",dest])
    print('success!')

## windows doc to docx
from win32com import client
word = client.Dispatch("Word.Application")# 打开word应用程序
for file in files:
    doc = word.Documents.Open(file) #打开word文件
    doc.SaveAs("{}x".format(file), 12)#另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close() #关闭原来word文件
word.Quit()
print("完成！")


file = "/Users/zlw/Documents/17-临床路径/国家发布的1010个临床路径汇总/2016年新发布的临床路径(1-524)/1小儿气管（支气管）异物临床路径.doc"
doc = docx.Document(file)
for num in range(len(doc.paragraphs)):
    print(doc.paragraphs[num].text)